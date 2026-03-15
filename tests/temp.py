"""
PDFLOW - PDF to DOCX Conversion Engine
Core conversion logic implementing four-phase architecture
"""

import os
import json
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict, Any, Tuple
import logging
from pathlib import Path

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(_name_)


class PDFParser:
    """
    PHASE A: PDF Parsing and Element Extraction
    Extracts text, images, and structural information with precise coordinates
    """
    
    def __init__(self, pdf_path: str):
        self.pdf_path = pdf_path
        self.doc = fitz.open(pdf_path)
        self.page_data = []
        
    def extract_all_pages(self) -> List[Dict[str, Any]]:
        """Extract all elements from all pages"""
        logger.info(f"Starting extraction for {len(self.doc)} pages")
        
        for page_num in range(len(self.doc)):
            page_info = self.extract_page(page_num)
            self.page_data.append(page_info)
            
        return self.page_data
    
    def extract_page(self, page_num: int) -> Dict[str, Any]:
        """Extract all elements from a single page"""
        page = self.doc[page_num]
        
        page_info = {
            'page_number': page_num + 1,
            'width': page.rect.width,
            'height': page.rect.height,
            'elements': []
        }
        
        # Extract text blocks with formatting
        text_blocks = page.get_text("dict")["blocks"]
        for block in text_blocks:
            if block['type'] == 0:  # Text block
                for line in block.get('lines', []):
                    for span in line.get('spans', []):
                        element = {
                            'type': 'text',
                            'bbox': span['bbox'],  # [x0, y0, x1, y1]
                            'text': span['text'],
                            'font': span['font'],
                            'font_size': span['size'],
                            'color': self._rgb_to_hex(span['color']),
                            'flags': span['flags'],  # Bit flags for bold, italic
                            'bold': bool(span['flags'] & 2**4),
                            'italic': bool(span['flags'] & 2**1),
                        }
                        page_info['elements'].append(element)
        
        # Extract images
        image_list = page.get_images(full=True)
        for img_index, img in enumerate(image_list):
            xref = img[0]
            base_image = self.doc.extract_image(xref)
            
            # Get image position (approximate from page layout)
            img_rects = page.get_image_rects(xref)
            
            for rect in img_rects:
                element = {
                    'type': 'image',
                    'bbox': [rect.x0, rect.y0, rect.x1, rect.y1],
                    'image_data': base_image['image'],
                    'ext': base_image['ext'],
                    'width': base_image['width'],
                    'height': base_image['height'],
                    'xref': xref
                }
                page_info['elements'].append(element)
        
        # Detect tables (basic grid detection)
        tables = self._detect_tables(page_info['elements'])
        page_info['tables'] = tables
        
        return page_info
    
    def _rgb_to_hex(self, rgb_int: int) -> str:
        """Convert RGB integer to hex color"""
        r = (rgb_int >> 16) & 0xFF
        g = (rgb_int >> 😎 & 0xFF
        b = rgb_int & 0xFF
        return f"#{r:02x}{g:02x}{b:02x}"
    
    def _detect_tables(self, elements: List[Dict]) -> List[Dict]:
        """Basic table detection using alignment patterns"""
        # Simplified version - production would use more sophisticated algorithms
        tables = []
        
        # Group elements by Y-coordinate (rows)
        text_elements = [e for e in elements if e['type'] == 'text']
        if len(text_elements) < 4:  # Minimum for a 2x2 table
            return tables
        
        # Sort by Y position
        sorted_by_y = sorted(text_elements, key=lambda x: x['bbox'][1])
        
        # Detect regular spacing (table indicator)
        y_positions = [e['bbox'][1] for e in sorted_by_y]
        
        # Calculate gaps between consecutive elements
        gaps = [y_positions[i+1] - y_positions[i] for i in range(len(y_positions)-1)]
        
        # If we have consistent gaps, it might be a table
        if len(gaps) > 3:
            avg_gap = sum(gaps) / len(gaps)
            consistent_gaps = sum(1 for g in gaps if abs(g - avg_gap) < 5)
            
            if consistent_gaps / len(gaps) > 0.7:  # 70% consistency
                # This is likely a table - more sophisticated detection needed
                tables.append({
                    'bbox': [
                        min(e['bbox'][0] for e in text_elements),
                        min(e['bbox'][1] for e in text_elements),
                        max(e['bbox'][2] for e in text_elements),
                        max(e['bbox'][3] for e in text_elements)
                    ],
                    'elements': text_elements
                })
        
        return tables


class OCRProcessor:
    """
    PHASE B: OCR Processing for scanned content and image-based text
    Supports Nepali and multi-language recognition
    """
    
    def __init__(self, languages: List[str] = ['eng', 'nep']):
        self.languages = languages
        self.preeti_unicode_map = self._load_preeti_map()
        
    def _load_preeti_map(self) -> Dict[str, str]:
        """Load Preeti to Unicode conversion map"""
        # Simplified map - production version would have complete mapping
        return {
            'k': 'क',
            'K': 'ख',
            'g': 'ग',
            'G': 'घ',
            'c': 'च',
            'C': 'छ',
            'j': 'ज',
            'J': 'झ',
            # ... complete mapping would be loaded from file
        }
    
    def process_element(self, element: Dict[str, Any]) -> Dict[str, Any]:
        """Process an element that might need OCR"""
        
        if element['type'] == 'image':
            # Convert to PIL Image
            from io import BytesIO
            img = Image.open(BytesIO(element['image_data']))
            
            # Pre-process image
            img = self._preprocess_image(img)
            
            # Run OCR
            lang_string = '+'.join(self.languages)
            
            try:
                # Tesseract with custom config
                custom_config = r'--oem 3 --psm 6'  # LSTM engine, assume uniform block
                text = pytesseract.image_to_string(
                    img, 
                    lang=lang_string,
                    config=custom_config
                )
                
                # Detect if this is Preeti font (heuristic)
                if self._is_preeti_encoded(text):
                    text = self._convert_preeti_to_unicode(text)
                
                element['ocr_text'] = text
                element['ocr_confidence'] = self._get_confidence(img, lang_string)
                
            except Exception as e:
                logger.error(f"OCR failed: {str(e)}")
                element['ocr_text'] = ""
                element['ocr_confidence'] = 0.0
        
        return element
    
    def _preprocess_image(self, img: Image.Image) -> Image.Image:
        """Pre-process image for better OCR accuracy"""
        # Convert to grayscale
        img = img.convert('L')
        
        # Enhance contrast
        from PIL import ImageEnhance
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(2.0)
        
        # Denoise (simple threshold)
        import numpy as np
        img_array = np.array(img)
        img_array[img_array < 128] = 0
        img_array[img_array >= 128] = 255
        img = Image.fromarray(img_array)
        
        return img
    
    def _is_preeti_encoded(self, text: str) -> bool:
        """Detect if text is Preeti-encoded (heuristic)"""
        # Check for common Preeti characters
        preeti_chars = set('kKgGcCjJ')
        text_chars = set(text)
        
        # If more than 30% are Preeti chars, assume Preeti
        if len(text_chars) > 0:
            overlap = len(preeti_chars & text_chars)
            return overlap / len(text_chars) > 0.3
        
        return False
    
    def _convert_preeti_to_unicode(self, preeti_text: str) -> str:
        """Convert Preeti-encoded text to Unicode Devanagari"""
        import unicodedata
        
        result = []
        for char in preeti_text:
            unicode_char = self.preeti_unicode_map.get(char, char)
            result.append(unicode_char)
        
        # Normalize to NFC (canonical composition)
        return unicodedata.normalize('NFC', ''.join(result))
    
    def _get_confidence(self, img: Image.Image, lang: str) -> float:
        """Get OCR confidence score"""
        try:
            data = pytesseract.image_to_data(img, lang=lang, output_type=pytesseract.Output.DICT)
            confidences = [int(conf) for conf in data['conf'] if conf != '-1']
            return sum(confidences) / len(confidences) if confidences else 0.0
        except:
            return 0.0


class LayoutAnalyzer:
    """
    PHASE C: Structural Layout Analysis
    Analyzes document structure to prevent text-box fragmentation
    """
    
    def __init__(self):
        self.tolerance_vertical = 5  # pixels
        self.tolerance_horizontal = 10
        
    def analyze_structure(self, page_data: Dict[str, Any]) -> Dict[str, Any]:
        """Analyze page structure and group elements semantically"""
        
        elements = page_data['elements']
        text_elements = [e for e in elements if e['type'] == 'text']
        
        # Step 1: Cluster into lines
        lines = self._cluster_by_y(text_elements)
        
        # Step 2: Detect columns
        columns = self._detect_columns(lines, page_data['width'])
        
        # Step 3: Group into paragraphs
        paragraphs = self._group_paragraphs(lines)
        
        # Step 4: Detect headers/footers
        headers = self._detect_headers(text_elements, page_data['height'])
        footers = self._detect_footers(text_elements, page_data['height'])
        
        # Step 5: Detect lists
        lists = self._detect_lists(paragraphs)
        
        return {
            'lines': lines,
            'columns': columns,
            'paragraphs': paragraphs,
            'headers': headers,
            'footers': footers,
            'lists': lists
        }
    
    def _cluster_by_y(self, elements: List[Dict]) -> List[List[Dict]]:
        """Cluster elements into lines by Y-coordinate"""
        if not elements:
            return []
        
        sorted_elements = sorted(elements, key=lambda x: x['bbox'][1])
        lines = []
        current_line = [sorted_elements[0]]
        
        for elem in sorted_elements[1:]:
            # If Y-coordinate is close to previous, same line
            if abs(elem['bbox'][1] - current_line[-1]['bbox'][1]) < self.tolerance_vertical:
                current_line.append(elem)
            else:
                lines.append(sorted(current_line, key=lambda x: x['bbox'][0]))
                current_line = [elem]
        
        if current_line:
            lines.append(sorted(current_line, key=lambda x: x['bbox'][0]))
        
        return lines
    
    def _detect_columns(self, lines: List[List[Dict]], page_width: float) -> int:
        """Detect number of columns in the page"""
        if not lines:
            return 1
        
        # Sample mid-section of page to avoid headers/footers
        mid_start = len(lines) // 3
        mid_end = 2 * len(lines) // 3
        mid_lines = lines[mid_start:mid_end]
        
        # Find X-position gaps
        all_x_positions = set()
        for line in mid_lines:
            for elem in line:
                all_x_positions.add(elem['bbox'][0])
        
        sorted_x = sorted(all_x_positions)
        
        # Detect large gaps (column separators)
        gaps = []
        for i in range(len(sorted_x) - 1):
            gap = sorted_x[i+1] - sorted_x[i]
            if gap > page_width / 10:  # Gap larger than 10% of page width
                gaps.append((sorted_x[i], gap))
        
        return len(gaps) + 1
    
    def _group_paragraphs(self, lines: List[List[Dict]]) -> List[Dict]:
        """Group lines into paragraphs"""
        paragraphs = []
        current_para_lines = []
        last_x_start = None
        
        for line in lines:
            if not line:
                continue
            
            x_start = line[0]['bbox'][0]
            
            # New paragraph if X-position changes significantly or large gap
            if last_x_start is not None:
                if abs(x_start - last_x_start) > self.tolerance_horizontal:
                    # Save current paragraph
                    if current_para_lines:
                        paragraphs.append(self._create_paragraph(current_para_lines))
                        current_para_lines = []
            
            current_para_lines.append(line)
            last_x_start = x_start
        
        # Add last paragraph
        if current_para_lines:
            paragraphs.append(self._create_paragraph(current_para_lines))
        
        return paragraphs
    
    def _create_paragraph(self, lines: List[List[Dict]]) -> Dict:
        """Create paragraph structure from lines"""
        all_elements = [elem for line in lines for elem in line]
        
        # Calculate bounding box
        bbox = [
            min(e['bbox'][0] for e in all_elements),
            min(e['bbox'][1] for e in all_elements),
            max(e['bbox'][2] for e in all_elements),
            max(e['bbox'][3] for e in all_elements)
        ]
        
        # Combine text
        text = ' '.join(' '.join(e['text'] for e in line) for line in lines)
        
        # Get predominant formatting
        font = max(set(e['font'] for e in all_elements), 
                  key=lambda x: sum(1 for e in all_elements if e['font'] == x))
        font_size = sum(e['font_size'] for e in all_elements) / len(all_elements)
        
        # Calculate indentation
        first_line_x = lines[0][0]['bbox'][0] if lines and lines[0] else 0
        subsequent_x = lines[1][0]['bbox'][0] if len(lines) > 1 and lines[1] else first_line_x
        
        return {
            'type': 'paragraph',
            'bbox': bbox,
            'text': text,
            'font': font,
            'font_size': font_size,
            'indent_first_line': first_line_x,
            'indent_left': subsequent_x,
            'lines': lines
        }
    
    def _detect_headers(self, elements: List[Dict], page_height: float) -> List[Dict]:
        """Detect header elements (top 15% of page)"""
        header_threshold = page_height * 0.15
        return [e for e in elements if e['bbox'][1] < header_threshold]
    
    def _detect_footers(self, elements: List[Dict], page_height: float) -> List[Dict]:
        """Detect footer elements (bottom 10% of page)"""
        footer_threshold = page_height * 0.90
        return [e for e in elements if e['bbox'][1] > footer_threshold]
    
    def _detect_lists(self, paragraphs: List[Dict]) -> List[Dict]:
        """Detect list structures"""
        lists = []
        
        # Look for bullet characters or numbering
        bullet_chars = {'•', '◦', '▪', '▫', '-', '*'}
        
        for para in paragraphs:
            text = para['text'].strip()
            
            # Check for bullet
            if text and text[0] in bullet_chars:
                para['list_type'] = 'bullet'
                para['list_level'] = self._calculate_list_level(para['indent_left'])
                lists.append(para)
            
            # Check for numbering (1., 2., a., etc.)
            elif text and len(text) > 2:
                if text[0].isdigit() or text[0].isalpha():
                    if text[1] in '.):':
                        para['list_type'] = 'numbered'
                        para['list_level'] = self._calculate_list_level(para['indent_left'])
                        lists.append(para)
        
        return lists
    
    def _calculate_list_level(self, indent: float) -> int:
        """Calculate list nesting level from indentation"""
        # Assuming 0.5 inch (36 points) per level
        return int(indent / 36)


class DOCXGenerator:
    """
    PHASE D: DOCX Document Generation
    Creates native OpenXML .docx files with preserved layout
    """
    
    def __init__(self):
        self.font_fallback_map = {
            'Preeti': 'Mangal',
            'Kantipur': 'Kokila',
            'Arial': 'Arial',
            'Times': 'Times New Roman',
            'Helvetica': 'Arial',
            'DEFAULT': 'Mangal'
        }
        
    def generate_docx(self, page_data_list: List[Dict], 
                     structure_list: List[Dict], 
                     output_path: str):
        """Generate DOCX from analyzed page data"""
        
        doc = Document()
        
        # Set page size (A4 by default)
        section = doc.sections[0]
        section.page_height = Inches(11.69)
        section.page_width = Inches(8.27)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Process each page
        for page_idx, (page_data, structure) in enumerate(zip(page_data_list, structure_list)):
            self._add_page_content(doc, page_data, structure)
            
            # Add page break (except for last page)
            if page_idx < len(page_data_list) - 1:
                doc.add_page_break()
        
        # Save document
        doc.save(output_path)
        logger.info(f"DOCX saved to {output_path}")
        
        return output_path
    
    def _add_page_content(self, doc: Document, 
                         page_data: Dict, 
                         structure: Dict):
        """Add content from a single page"""
        
        # Add headers first
        if structure.get('headers'):
            for header_elem in structure['headers']:
                self._add_text_element(doc, header_elem, is_header=True)
        
        # Add paragraphs
        for para in structure.get('paragraphs', []):
            self._add_paragraph(doc, para)
        
        # Add images
        for elem in page_data['elements']:
            if elem['type'] == 'image':
                self._add_image(doc, elem)
        
        # Add footers
        if structure.get('footers'):
            # Note: Actual footer implementation requires section.footer
            pass
    
    def _add_paragraph(self, doc: Document, para_data: Dict):
        """Add a paragraph with formatting"""
        
        p = doc.add_paragraph()
        
        # Set indentation (convert points to inches)
        indent_first = para_data.get('indent_first_line', 0) / 72
        indent_left = para_data.get('indent_left', 0) / 72
        
        p.paragraph_format.first_line_indent = Inches(indent_first)
        p.paragraph_format.left_indent = Inches(indent_left)
        
        # Handle list formatting
        if para_data.get('list_type'):
            # For production, use proper numbering definitions
            # This is simplified
            bullet = '• ' if para_data['list_type'] == 'bullet' else '1. '
            text = bullet + para_data['text']
        else:
            text = para_data['text']
        
        # Add text with formatting
        run = p.add_run(text)
        
        # Map font
        font_name = para_data.get('font', 'DEFAULT')
        mapped_font = self.font_fallback_map.get(font_name, self.font_fallback_map['DEFAULT'])
        run.font.name = mapped_font
        
        # Set font size
        font_size = para_data.get('font_size', 12)
        run.font.size = Pt(font_size)
        
        # Set color if available
        color_hex = para_data.get('color', '#000000')
        if color_hex != '#000000':
            r, g, b = self._hex_to_rgb(color_hex)
            run.font.color.rgb = RGBColor(r, g, b)
    
    def _add_text_element(self, doc: Document, element: Dict, is_header: bool = False):
        """Add a single text element"""
        p = doc.add_paragraph()
        
        if is_header:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = p.add_run(element['text'])
        
        # Font mapping
        font = self.font_fallback_map.get(element.get('font'), self.font_fallback_map['DEFAULT'])
        run.font.name = font
        run.font.size = Pt(element.get('font_size', 12))
        
        # Bold/Italic
        run.bold = element.get('bold', False)
        run.italic = element.get('italic', False)
    
    def _add_image(self, doc: Document, image_elem: Dict):
        """Add an image to the document"""
        from io import BytesIO
        
        try:
            # Create image from bytes
            image_stream = BytesIO(image_elem['image_data'])
            
            # Calculate size (convert from pixels to inches, assuming 72 DPI)
            width_inches = image_elem['width'] / 72
            height_inches = image_elem['height'] / 72
            
            # Add to document
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_picture(image_stream, width=Inches(width_inches))
            
        except Exception as e:
            logger.error(f"Failed to add image: {str(e)}")
    
    def _hex_to_rgb(self, hex_color: str) -> Tuple[int, int, int]:
        """Convert hex color to RGB tuple"""
        hex_color = hex_color.lstrip('#')
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


class PDFToDOCXConverter:
    """
    Main converter class orchestrating all four phases
    """
    
    def __init__(self, languages: List[str] = ['eng', 'nep']):
        self.languages = languages
        self.parser = None
        self.ocr_processor = OCRProcessor(languages)
        self.layout_analyzer = LayoutAnalyzer()
        self.docx_generator = DOCXGenerator()
    
    def convert(self, pdf_path: str, output_path: str = None) -> str:
        """
        Convert PDF to DOCX with full pipeline
        
        Args:
            pdf_path: Path to input PDF
            output_path: Path for output DOCX (optional)
            
        Returns:
            Path to generated DOCX file
        """
        
        if output_path is None:
            output_path = pdf_path.rsplit('.', 1)[0] + '.docx'
        
        logger.info(f"Starting conversion: {pdf_path} -> {output_path}")
        
        # PHASE A: Parse PDF
        logger.info("Phase A: Parsing PDF...")
        self.parser = PDFParser(pdf_path)
        page_data_list = self.parser.extract_all_pages()
        
        # PHASE B: OCR Processing (if needed)
        logger.info("Phase B: OCR Processing...")
        for page_data in page_data_list:
            for elem in page_data['elements']:
                if elem['type'] == 'image':
                    self.ocr_processor.process_element(elem)
        
        # PHASE C: Layout Analysis
        logger.info("Phase C: Layout Analysis...")
        structure_list = []
        for page_data in page_data_list:
            structure = self.layout_analyzer.analyze_structure(page_data)
            structure_list.append(structure)
        
        # PHASE D: DOCX Generation
        logger.info("Phase D: Generating DOCX...")
        output = self.docx_generator.generate_docx(page_data_list, structure_list, output_path)
        
        logger.info(f"Conversion complete: {output}")
        return output


# Example usage
if _name_ == "_main_":
    converter = PDFToDOCXConverter(languages=['eng', 'nep'])
    
    # Convert a sample PDF
    result = converter.convert(
        pdf_path="sample.pdf",
        output_path="output.docx"
    )
    
    print(f"Conversion successful: {result}")